#!/usr/bin/env python3
"""Build Kindle files for the two Star of the Toy Farm stories.

Produces, for each title:
  *-kindle.docx     ← upload this to KDP first
  *-kindle.epub     ← EPUB 3 fallback
  kindle-cover.jpg

Source text is copied from the site HTML without rewriting Anshika's story.

    python3 kdp/star-of-the-toy-farm/build_kindle.py
"""
from __future__ import annotations

import html as html_lib
import io
import re
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageEnhance, ImageFont

ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
SITE = ROOT / "anshika-mahesh-site/stories"
FONT_DIR = Path("/System/Library/Fonts/Supplemental")

AUTHOR = "Anshika Mahesh"
SERIES = "Star of the Toy Farm"
LANG = "en-GB"
COVER_W, COVER_H = 1600, 2560
PLATE_MAX_W = 1200
INK = RGBColor(0x1F, 0x2A, 0x24)
TEAL = "3D7F99"

ABOUT_AUTHOR = (
    "Anshika Mahesh is 12 years old. She writes stories about friendship, "
    "courage, and kindness — from a kid, to a kid."
)

BOOKS = [
    {
        "id": "home",
        "out_dir": "star-finds-a-home",
        "slug": "star-finds-a-home",
        "html": SITE / "star-of-the-toy-farm.html",
        "cover_src": SITE / "images/star-of-the-toy-farm/cover.png",
        "title": "Star Finds a Home",
        "title_lines": ["STAR FINDS", "A HOME"],
        "subtitle": "A Star of the Toy Farm story",
        "blurb": (
            "Star is a toy sheep who wonders why she exists. When she becomes "
            "Sophia’s birthday gift, she finds friends on the toy farm — but a "
            "missing cow named Mia pulls her into her first real adventure."
        ),
        "dedication": "From a kid, to a kid.",
        "end_line": "And that is the best thing I could ever own.",
        "uuid_url": "https://anshikamahesh.com/stories/star-of-the-toy-farm.html",
        "next_note": "The next Star of the Toy Farm story is Star and the Midnight Feast.",
        "overlay": (111, 140, 90, 168),
        "expected_chapters": 12,
    },
    {
        "id": "feast",
        "out_dir": "star-and-the-midnight-feast",
        "slug": "star-and-the-midnight-feast",
        "html": SITE / "star-and-the-midnight-feast.html",
        "cover_src": SITE / "images/star-of-the-toy-farm/midnight-feast.png",
        "title": "Star and the Midnight Feast",
        "title_lines": ["STAR AND THE", "MIDNIGHT FEAST"],
        "subtitle": "A Star of the Toy Farm story",
        "blurb": (
            "A sleepover, a show-and-tell prize, a riddle-loving goat, and a "
            "midnight feast — Star learns that bravery isn’t only mountains. "
            "It’s looking after friends when the night gets messy."
        ),
        "dedication": "Best after Star Finds a Home.",
        "end_line": "I am wondering what adventure comes next.",
        "uuid_url": "https://anshikamahesh.com/stories/star-and-the-midnight-feast.html",
        "next_note": "",
        "overlay": (61, 127, 153, 168),
        "expected_chapters": 12,
    },
]


def esc(text: str) -> str:
    return html_lib.escape(text, quote=True)


def clean(frag: str) -> str:
    t = re.sub(r"<br\s*/?>", " ", frag, flags=re.I)
    t = re.sub(r"<[^>]+>", "", t)
    t = html_lib.unescape(t)
    return re.sub(r"\s+", " ", t).strip()


def jpeg_bytes(src: Path, max_w: int, quality: int = 82) -> bytes:
    im = Image.open(src).convert("RGB")
    if im.width > max_w:
        h = max(1, int(im.height * max_w / im.width))
        im = im.resize((max_w, h), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, "JPEG", quality=quality, optimize=True, progressive=False)
    return buf.getvalue()


def parse_figure(fig_html: str, story_html: Path) -> dict | None:
    src_m = re.search(r'<img[^>]+src="([^"]+)"', fig_html)
    if not src_m:
        return None
    src = (story_html.parent / src_m.group(1).split("?")[0]).resolve()
    alt_m = re.search(r'alt="([^"]*)"', fig_html)
    cap_m = re.search(r"<figcaption>(.*?)</figcaption>", fig_html, re.S)
    return {
        "type": "image",
        "path": src,
        "alt": html_lib.unescape(alt_m.group(1)) if alt_m else "",
        "caption": clean(cap_m.group(1)) if cap_m else "",
    }


def parse_script(script_html: str) -> list[dict]:
    blocks = []
    for p in re.findall(r"<p[^>]*>(.*?)</p>", script_html, re.S):
        m = re.match(r"<strong>(.*?)</strong>\s*(.*)", p, re.S)
        if not m:
            text = clean(p)
            if text:
                blocks.append({"type": "p", "text": text, "kind": ""})
            continue
        speaker = clean(m.group(1)).replace("—", "").strip(" -")
        rest = m.group(2)
        stage_m = re.search(r'<span class="stage">(.*?)</span>', rest, re.S)
        dialogue = clean(re.sub(r'<span class="stage">.*?</span>', " ", rest, flags=re.S))
        blocks.append(
            {
                "type": "line",
                "speaker": speaker,
                "text": dialogue,
                "stage": clean(stage_m.group(1)) if stage_m else "",
            }
        )
    return blocks


def parse_playbox(box_html: str) -> list[dict]:
    blocks = []
    h = re.search(r"<h3>(.*?)</h3>", box_html, re.S)
    if h:
        blocks.append({"type": "box_title", "text": clean(h.group(1))})
    for item in re.findall(r"<li>(.*?)</li>", box_html, re.S):
        text = clean(item)
        if text:
            blocks.append({"type": "bullet", "text": text})
    return blocks


def parse_blocks(body: str, story_html: Path) -> list[dict]:
    token = re.compile(
        r'(<div class="gallery">.*?</div>)'
        r'|(<figure class="illust">.*?</figure>)'
        r'|(<div class="script">.*?</div>)'
        r'|(<div class="playbox">.*?</div>)'
        r'|(<div class="chant">.*?</div>)'
        r'|(<div class="nick">.*?</div>)'
        r'|(<ul class="tips">.*?</ul>)'
        r'|(<p\b[^>]*>.*?</p>)',
        re.S,
    )
    blocks: list[dict] = []
    for gallery, figure, script, playbox, chant, nick, tips, para in token.findall(body):
        if gallery:
            for fig in re.findall(r'<figure class="illust">.*?</figure>', gallery, re.S):
                parsed = parse_figure(fig, story_html)
                if parsed and parsed["path"].exists():
                    blocks.append(parsed)
            continue
        if figure:
            parsed = parse_figure(figure, story_html)
            if parsed and parsed["path"].exists():
                blocks.append(parsed)
            continue
        if script:
            blocks.extend(parse_script(script))
            continue
        if playbox:
            blocks.extend(parse_playbox(playbox))
            continue
        if chant:
            for p in re.findall(r"<p[^>]*>(.*?)</p>", chant, re.S):
                text = clean(p)
                if text:
                    blocks.append({"type": "chant", "text": text})
            continue
        if nick:
            for line in re.split(r"<br\s*/?>", nick, flags=re.I):
                text = clean(line)
                if text:
                    blocks.append({"type": "bullet", "text": text})
            continue
        if tips:
            for item in re.findall(r"<li>(.*?)</li>", tips, re.S):
                text = clean(item)
                if text:
                    blocks.append({"type": "bullet", "text": text})
            continue
        inner_m = re.search(r"<p\b([^>]*)>(.*?)</p>", para, re.S)
        if not inner_m:
            continue
        attrs, inner = inner_m.group(1), inner_m.group(2)
        if re.search(r'<a\s+href=', inner, re.I):
            continue
        text = clean(inner)
        if not text:
            continue
        if "end-mark" in attrs:
            continue
        if "pov" in attrs:
            blocks.append({"type": "pov", "text": text})
        elif "thought" in attrs:
            blocks.append({"type": "thought", "text": text})
        elif "shout" in attrs:
            blocks.append({"type": "shout", "text": text})
        elif "stage" in attrs:
            blocks.append({"type": "stage", "text": text})
        elif "note" in attrs:
            blocks.append({"type": "p", "text": text, "kind": "dropcap"})
        else:
            blocks.append({"type": "p", "text": text, "kind": "dropcap" if "dropcap" in attrs else ""})
    return blocks


def parse_cast(raw: str) -> list[dict]:
    cast = []
    section = re.search(r'<section class="page" id="cast">(.*?)</section>', raw, re.S)
    if not section:
        return cast
    for row in re.findall(r"<div><strong>(.*?)</strong>\s*[—\-]\s*(.*?)</div>", section.group(1), re.S):
        cast.append({"name": clean(row[0]), "text": clean(row[1])})
    return cast


def parse_story(book: dict) -> dict:
    raw = book["html"].read_text(encoding="utf-8")
    chapters = []
    for num, body in re.findall(r'<section class="page" id="ch(\d+)">(.*?)</section>', raw, re.S):
        title_m = re.search(r"<h2>([^<]+)</h2>", body)
        header, _, rest = body.partition("</header>")
        chapter_body = rest if rest else body
        chapters.append(
            {
                "num": int(num),
                "title": html_lib.unescape(title_m.group(1).strip()) if title_m else f"Chapter {num}",
                "blocks": parse_blocks(chapter_body, book["html"]),
            }
        )
    end_img = None
    end_sec = re.search(r'<section class="page end-page">(.*?)</section>', raw, re.S)
    if end_sec:
        fig = re.search(r'<figure class="illust">.*?</figure>', end_sec.group(1), re.S)
        if fig:
            end_img = parse_figure(fig.group(0), book["html"])
    return {"chapters": chapters, "cast": parse_cast(raw), "end_image": end_img}


def build_cover_jpeg(book: dict, out_path: Path) -> None:
    img = Image.new("RGB", (COVER_W, COVER_H), "#1f2a24")
    art = ImageEnhance.Contrast(Image.open(book["cover_src"]).convert("RGB")).enhance(1.02)
    scale = max(COVER_W / art.width, COVER_H / art.height)
    nw, nh = int(art.width * scale), int(art.height * scale)
    art = art.resize((nw, nh), Image.Resampling.LANCZOS)
    img.paste(art, ((COVER_W - nw) // 2, (COVER_H - nh) // 2))

    overlay = Image.new("RGBA", (COVER_W, COVER_H), (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)
    od.rectangle((0, 0, COVER_W, 470), fill=book["overlay"])
    od.rectangle((0, COVER_H - 260, COVER_W, COVER_H), fill=book["overlay"])
    img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype(str(FONT_DIR / "Georgia Bold.ttf"), 72)
        sub_font = ImageFont.truetype(str(FONT_DIR / "Georgia Italic.ttf"), 28)
        author_font = ImageFont.truetype(str(FONT_DIR / "Georgia Bold.ttf"), 32)
    except OSError:
        title_font = sub_font = author_font = ImageFont.load_default()

    def centre(text: str, font, y: float, fill="white") -> None:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(((COVER_W - tw) / 2, y - th / 2), text, font=font, fill=fill)

    centre(book["title_lines"][0], title_font, 150)
    centre(book["title_lines"][1], title_font, 250)
    centre(book["subtitle"], sub_font, 355)
    centre("ANSHIKA MAHESH", author_font, COVER_H - 130)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, "JPEG", quality=90, optimize=True, progressive=False, dpi=(300, 300))


def styles_css() -> str:
    return """body {
  font-family: Georgia, "Times New Roman", serif;
  font-size: 1em;
  line-height: 1.45;
  margin: 0;
  text-align: justify;
}
h1 {
  font-size: 1.4em;
  text-align: center;
  margin: 0.4em 0 0.8em 0;
  page-break-after: avoid;
}
.kicker, .pov, .stage, .caption, .center { text-align: center; text-indent: 0; }
.kicker, .pov { font-size: 0.85em; font-weight: bold; text-transform: uppercase; margin: 1.2em 0 0.4em 0; }
.thought, .shout, .chant, .stage { font-style: italic; text-indent: 0; }
.shout { text-align: center; }
p { margin: 0 0 0.6em 0; text-indent: 1.2em; }
p.first, p.line { text-indent: 0; }
img.plate { width: 100%; height: auto; display: block; margin: 0.5em 0 0.3em 0; }
.caption { font-size: 0.85em; font-style: italic; margin: 0 0 1em 0; }
"""


def container_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
"""


def xhtml(title: str, body: str) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" '
        '"http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">\n'
        f'<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{LANG}">\n'
        "<head>\n"
        f"<title>{esc(title)}</title>\n"
        '<meta http-equiv="Content-Type" content="application/xhtml+xml; charset=utf-8" />\n'
        '<link rel="stylesheet" type="text/css" href="styles.css" />\n'
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )


def block_xhtml(block: dict, img_name: str | None) -> str:
    t = block["type"]
    if t == "image" and img_name:
        cap = f'\n<p class="caption">{esc(block["caption"])}</p>' if block.get("caption") else ""
        return f'<img class="plate" src="images/{img_name}" alt="{esc(block.get("alt", ""))}" />{cap}'
    if t == "pov":
        return f'<p class="pov">{esc(block["text"])}</p>'
    if t == "thought":
        return f'<p class="thought first">{esc(block["text"])}</p>'
    if t == "shout":
        return f'<p class="shout first">{esc(block["text"])}</p>'
    if t == "stage":
        return f'<p class="stage first">{esc(block["text"])}</p>'
    if t == "chant":
        return f'<p class="chant first">{esc(block["text"])}</p>'
    if t == "box_title":
        return f'<p class="kicker">{esc(block["text"])}</p>'
    if t == "bullet":
        return f'<p class="first">• {esc(block["text"])}</p>'
    if t == "line":
        stage = f' <em>{esc(block["stage"])}</em>' if block.get("stage") else ""
        spoken = esc(block["text"])
        return f'<p class="line"><strong>{esc(block["speaker"])} —</strong> {spoken}{stage}</p>'
    cls = "first" if block.get("kind") == "dropcap" else ""
    attr = f' class="{cls}"' if cls else ""
    return f"<p{attr}>{esc(block['text'])}</p>"


def chapter_images(ch: dict) -> list[dict]:
    return [b for b in ch["blocks"] if b["type"] == "image"]


def content_opf(book: dict, data: dict, image_files: list[str]) -> str:
    items = [
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
        '<item id="css" href="styles.css" media-type="text/css"/>',
        '<item id="cover-image" href="images/cover.jpg" media-type="image/jpeg"/>',
        '<item id="titlepage" href="title.xhtml" media-type="application/xhtml+xml"/>',
        '<item id="copyright" href="copyright.xhtml" media-type="application/xhtml+xml"/>',
        '<item id="contents" href="contents.xhtml" media-type="application/xhtml+xml"/>',
    ]
    spine = [
        '<itemref idref="titlepage"/>',
        '<itemref idref="copyright"/>',
        '<itemref idref="contents"/>',
    ]
    if data["cast"]:
        items.append('<item id="cast" href="cast.xhtml" media-type="application/xhtml+xml"/>')
        spine.insert(2, '<itemref idref="cast"/>')
    for ch in data["chapters"]:
        n = f"{ch['num']:02d}"
        items.append(f'<item id="ch{n}" href="ch{n}.xhtml" media-type="application/xhtml+xml"/>')
        spine.append(f'<itemref idref="ch{n}"/>')
    items.append('<item id="theend" href="end.xhtml" media-type="application/xhtml+xml"/>')
    items.append('<item id="about" href="about.xhtml" media-type="application/xhtml+xml"/>')
    spine.append('<itemref idref="theend"/>')
    spine.append('<itemref idref="about"/>')
    for i, name in enumerate(image_files, 1):
        items.append(f'<item id="img{i:02d}" href="images/{name}" media-type="image/jpeg"/>')
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
    <dc:identifier id="BookId">urn:uuid:{book["book_id"]}</dc:identifier>
    <dc:title>{esc(book["title"])}</dc:title>
    <dc:creator>{esc(AUTHOR)}</dc:creator>
    <dc:language>{LANG}</dc:language>
    <dc:publisher>{esc(AUTHOR)}</dc:publisher>
    <dc:rights>Copyright &#169; 2026 {esc(AUTHOR)}. All rights reserved.</dc:rights>
    <dc:description>{esc(book["blurb"])}</dc:description>
    <dc:subject>{esc(SERIES)}</dc:subject>
    <meta property="dcterms:modified">2026-08-17T00:00:00Z</meta>
    <meta name="cover" content="cover-image"/>
  </metadata>
  <manifest>
    {chr(10).join("    " + i for i in items)}
  </manifest>
  <spine toc="ncx">
    {chr(10).join("    " + s for s in spine)}
  </spine>
  <guide>
    <reference type="cover" title="Cover" href="images/cover.jpg"/>
    <reference type="toc" title="Contents" href="contents.xhtml"/>
    <reference type="text" title="Start Reading" href="ch01.xhtml"/>
  </guide>
</package>
"""


def toc_ncx(book: dict, data: dict) -> str:
    points = [
        '<navPoint id="nav1" playOrder="1"><navLabel><text>Title</text></navLabel><content src="title.xhtml"/></navPoint>',
        '<navPoint id="nav2" playOrder="2"><navLabel><text>Contents</text></navLabel><content src="contents.xhtml"/></navPoint>',
    ]
    order = 3
    if data["cast"]:
        points.append(
            f'<navPoint id="nav-cast" playOrder="{order}"><navLabel><text>Main characters</text></navLabel><content src="cast.xhtml"/></navPoint>'
        )
        order += 1
    for ch in data["chapters"]:
        n = f"{ch['num']:02d}"
        points.append(
            f'<navPoint id="nav-ch{n}" playOrder="{order}">'
            f"<navLabel><text>Chapter {ch['num']}. {esc(ch['title'])}</text></navLabel>"
            f'<content src="ch{n}.xhtml"/></navPoint>'
        )
        order += 1
    points.append(
        f'<navPoint id="nav-end" playOrder="{order}"><navLabel><text>The End</text></navLabel><content src="end.xhtml"/></navPoint>'
    )
    points.append(
        f'<navPoint id="nav-about" playOrder="{order + 1}"><navLabel><text>About the Author</text></navLabel><content src="about.xhtml"/></navPoint>'
    )
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{book["book_id"]}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{esc(book["title"])}</text></docTitle>
  <navMap>
    {chr(10).join("    " + p for p in points)}
  </navMap>
</ncx>
"""


def nav_xhtml(data: dict) -> str:
    items = ['<li><a href="contents.xhtml">Contents</a></li>']
    if data["cast"]:
        items.append('<li><a href="cast.xhtml">Main characters</a></li>')
    for ch in data["chapters"]:
        items.append(
            f'<li><a href="ch{ch["num"]:02d}.xhtml">Chapter {ch["num"]}. {esc(ch["title"])}</a></li>'
        )
    items.append('<li><a href="end.xhtml">The End</a></li>')
    items.append('<li><a href="about.xhtml">About the Author</a></li>')
    body = (
        '<nav epub:type="toc" id="toc">\n'
        "<h1>Contents</h1>\n"
        "<ol>\n" + "\n".join(items) + "\n</ol>\n"
        "</nav>"
    )
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" '
        'xmlns:epub="http://www.idpf.org/2007/ops" '
        f'xml:lang="{LANG}">\n'
        "<head>\n<title>Contents</title>\n"
        '<link rel="stylesheet" type="text/css" href="styles.css" />\n'
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )


def _zipinfo(name: str, stored: bool = False) -> zipfile.ZipInfo:
    info = zipfile.ZipInfo(name, date_time=(2026, 8, 17, 0, 0, 0))
    info.compress_type = zipfile.ZIP_STORED if stored else zipfile.ZIP_DEFLATED
    info.create_system = 3
    info.external_attr = 0o644 << 16
    info.extra = b""
    return info


def assign_images(data: dict) -> dict[int, str]:
    """Map id(block) -> filename inside the epub."""
    names = {}
    n = 1
    for ch in data["chapters"]:
        for block in chapter_images(ch):
            names[id(block)] = f"img-{n:02d}.jpg"
            n += 1
    if data.get("end_image"):
        names[id(data["end_image"])] = f"img-{n:02d}.jpg"
    return names


def build_epub(book: dict, data: dict, cover_path: Path, epub_path: Path) -> int:
    img_names = assign_images(data)
    files: list[tuple[zipfile.ZipInfo, bytes]] = []

    def add(name: str, content: str | bytes, stored: bool = False) -> None:
        payload = content.encode("utf-8") if isinstance(content, str) else content
        files.append((_zipinfo(name, stored=stored), payload))

    add("mimetype", "application/epub+zip", stored=True)
    add("META-INF/container.xml", container_xml())
    add("OEBPS/styles.css", styles_css())
    add(
        "OEBPS/title.xhtml",
        xhtml(
            book["title"],
            f'<p class="kicker">{esc(SERIES)}</p>\n'
            f'<h1>{esc(book["title"])}</h1>\n'
            f'<p class="center">{esc(book["subtitle"])}</p>\n'
            f'<p class="center">{esc(AUTHOR)}</p>\n'
            f'<p class="first center"><em>{esc(book["dedication"])}</em></p>\n'
            f'<p class="first center">{esc(book["blurb"])}</p>',
        ),
    )
    add(
        "OEBPS/copyright.xhtml",
        xhtml(
            "Copyright",
            "<h1>Copyright</h1>\n"
            f'<p class="first"><strong>{esc(book["title"])}</strong></p>\n'
            f'<p class="first">A {esc(SERIES)} story</p>\n'
            f'<p class="first">Copyright &#169; 2026 {esc(AUTHOR)}</p>\n'
            f'<p class="first">First Kindle edition 2026</p>\n'
            f'<p class="first">All rights reserved. No part of this book may be reproduced without permission from the author, except for brief quotations in a review.</p>\n'
            f'<p class="first">This is a work of fiction. Names, characters, and incidents are from the author&#8217;s imagination.</p>',
        ),
    )
    if data["cast"]:
        rows = "\n".join(
            f'<p class="first"><strong>{esc(c["name"])}</strong> — {esc(c["text"])}</p>' for c in data["cast"]
        )
        add("OEBPS/cast.xhtml", xhtml("Main characters", "<h1>Main characters</h1>\n" + rows))
    toc_items = []
    if data["cast"]:
        toc_items.append('<p class="first"><a href="cast.xhtml">Main characters</a></p>')
    for ch in data["chapters"]:
        toc_items.append(
            f'<p class="first"><a href="ch{ch["num"]:02d}.xhtml">{ch["num"]}. {esc(ch["title"])}</a></p>'
        )
    add("OEBPS/contents.xhtml", xhtml("Contents", '<h1 id="toc">Contents</h1>\n' + "\n".join(toc_items)))

    image_files = []
    for ch in data["chapters"]:
        n = f"{ch['num']:02d}"
        parts = [
            f'<p class="kicker">Chapter {ch["num"]}</p>',
            f'<h1>{esc(ch["title"])}</h1>',
        ]
        first_p = True
        for block in ch["blocks"]:
            if block["type"] == "image":
                name = img_names[id(block)]
                if name not in image_files:
                    image_files.append(name)
                    add(f"OEBPS/images/{name}", jpeg_bytes(block["path"], PLATE_MAX_W))
                parts.append(block_xhtml(block, name))
                first_p = True
                continue
            html_block = block_xhtml(block, None)
            if block["type"] == "p" and first_p and ' class="' not in html_block:
                html_block = html_block.replace("<p>", '<p class="first">', 1)
                first_p = False
            elif block["type"] == "p":
                first_p = False
            parts.append(html_block)
        add(f"OEBPS/ch{n}.xhtml", xhtml(f"Chapter {ch['num']}. {ch['title']}", "\n".join(parts)))

    end_parts = ["<h1>The End</h1>", f'<p class="first center">{esc(book["end_line"])}</p>']
    if data.get("end_image"):
        name = img_names[id(data["end_image"])]
        image_files.append(name)
        add(f"OEBPS/images/{name}", jpeg_bytes(data["end_image"]["path"], PLATE_MAX_W))
        end_parts.append(block_xhtml(data["end_image"], name))
    add("OEBPS/end.xhtml", xhtml("The End", "\n".join(end_parts)))
    about = f'<p class="first">{esc(book["blurb"])}</p>\n<p>{esc(ABOUT_AUTHOR)}</p>'
    if book["next_note"]:
        about += f'\n<p>{esc(book["next_note"])}</p>'
    add("OEBPS/about.xhtml", xhtml("About the Author", "<h1>About the Author</h1>\n" + about))
    add("OEBPS/images/cover.jpg", cover_path.read_bytes())
    add("OEBPS/content.opf", content_opf(book, data, image_files))
    add("OEBPS/toc.ncx", toc_ncx(book, data))
    add("OEBPS/nav.xhtml", nav_xhtml(data))

    with zipfile.ZipFile(epub_path, "w") as zf:
        for info, payload in files:
            zf.writestr(info, payload)
    return epub_path.stat().st_size


def _set_run_font(run, name: str = "Georgia", size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = INK


def _center(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)


def _body(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15


def _bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _page_break(paragraph) -> None:
    paragraph.paragraph_format.page_break_before = True


def _add_toc_link(paragraph, anchor: str, label: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Georgia")
    fonts.set(qn("w:hAnsi"), "Georgia")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def _toc_field_run(kind: str, instr: str | None = None):
    r = OxmlElement("w:r")
    if kind == "instr":
        t = OxmlElement("w:instrText")
        t.set(qn("xml:space"), "preserve")
        t.text = instr or ""
        r.append(t)
        return r
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    if kind == "begin":
        fld.set(qn("w:dirty"), "true")
    r.append(fld)
    return r


def insert_kindle_toc(doc: Document, entries: list[tuple[str, str]], bookmark_id: int) -> int:
    """Word Automatic TOC with no page numbers, plus bookmark 'toc' — what KDP looks for."""
    title = doc.add_paragraph()
    try:
        title.style = doc.styles["TOC Heading"]
    except KeyError:
        pass
    _center(title)
    _page_break(title)
    run = title.add_run("Contents")
    _set_run_font(run, size=18, bold=True)
    _bookmark(title, "toc", bookmark_id)

    begin = doc.add_paragraph()
    begin.paragraph_format.first_line_indent = Inches(0)
    begin._p.append(_toc_field_run("begin"))
    begin._p.append(_toc_field_run("instr", r' TOC \o "1-1" \h \z \u '))
    begin._p.append(_toc_field_run("separate"))

    for anchor, label in entries:
        p = doc.add_paragraph()
        try:
            p.style = doc.styles["TOC 1"]
        except KeyError:
            pass
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(8)
        _add_toc_link(p, anchor, label)

    end = doc.add_paragraph()
    end.paragraph_format.first_line_indent = Inches(0)
    end._p.append(_toc_field_run("end"))
    return bookmark_id + 1


def add_picture(doc, path: Path, caption: str, width: float = 4.5) -> None:
    pic = doc.add_paragraph()
    _center(pic)
    pic.paragraph_format.space_after = Pt(4)
    run = pic.add_run()
    run.add_picture(io.BytesIO(jpeg_bytes(path, PLATE_MAX_W)), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        _center(cap)
        cap.paragraph_format.space_after = Pt(12)
        run = cap.add_run(caption)
        _set_run_font(run, size=10, italic=True)


def add_plain(doc, text: str, *, center=False, italic=False, bold=False, indent=0.0, size=12) -> None:
    p = doc.add_paragraph()
    if center:
        _center(p)
    else:
        _body(p)
        p.paragraph_format.first_line_indent = Inches(indent)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)


def write_block_docx(doc, block: dict, first_body: list[bool]) -> None:
    t = block["type"]
    if t == "image":
        wide = "costume" not in block["path"].name
        add_picture(doc, block["path"], block.get("caption", ""), width=4.5 if wide else 3.2)
        first_body[0] = True
        return
    if t == "pov":
        add_plain(doc, block["text"], center=True, bold=True, size=11)
        first_body[0] = True
        return
    if t == "thought":
        add_plain(doc, block["text"], italic=True, indent=0)
        return
    if t == "shout":
        add_plain(doc, block["text"], center=True, italic=True, bold=True)
        return
    if t == "stage":
        add_plain(doc, block["text"], center=True, italic=True, size=11)
        return
    if t == "chant":
        add_plain(doc, block["text"], italic=True, indent=0)
        return
    if t == "box_title":
        add_plain(doc, block["text"], center=True, bold=True, size=12)
        return
    if t == "bullet":
        add_plain(doc, "• " + block["text"], indent=0)
        return
    if t == "line":
        p = doc.add_paragraph()
        _body(p)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(f"{block['speaker']} — ")
        _set_run_font(run, size=12, bold=True)
        if block["text"]:
            run = p.add_run(block["text"])
            _set_run_font(run, size=12)
        if block.get("stage"):
            run = p.add_run(" " + block["stage"])
            _set_run_font(run, size=12, italic=True)
        return
    indent = 0 if first_body[0] or block.get("kind") == "dropcap" else 0.25
    add_plain(doc, block["text"], indent=indent)
    first_body[0] = False


def build_docx(book: dict, data: dict, docx_path: Path) -> None:
    doc = Document()
    core = doc.core_properties
    core.title = book["title"]
    core.author = AUTHOR
    core.subject = book["subtitle"]
    core.category = "Juvenile Fiction"

    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    heading = styles["Heading 1"]
    heading.font.name = "Georgia"
    heading.font.size = Pt(18)
    heading.font.bold = True
    heading.font.color.rgb = INK
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.page_break_before = True
    heading_ppr = heading.element.get_or_add_pPr()
    if heading_ppr.find(qn("w:outlineLvl")) is None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        heading_ppr.append(outline)
    _enable_update_fields(doc)

    add_plain(doc, SERIES, center=True, bold=True, size=11)
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run(book["title"])
    _set_run_font(run, size=26, bold=True)
    add_plain(doc, book["subtitle"], center=True, italic=True, size=13)
    add_plain(doc, AUTHOR, center=True, bold=True, size=12)
    add_plain(doc, book["dedication"], center=True, italic=True, size=12)
    add_plain(doc, book["blurb"], center=True, italic=True, size=12)

    p = doc.add_paragraph()
    _center(p)
    _page_break(p)
    run = p.add_run("Copyright")
    _set_run_font(run, size=18, bold=True)
    for line in (
        book["title"],
        f"A {SERIES} story",
        f"Copyright © 2026 {AUTHOR}",
        "First Kindle edition 2026",
        "All rights reserved. No part of this book may be reproduced without permission from the author, except for brief quotations in a review.",
        "This is a work of fiction. Names, characters, and incidents are from the author’s imagination.",
    ):
        add_plain(doc, line, indent=0, size=11)

    bookmark_id = 1
    toc_entries: list[tuple[str, str]] = []
    if data["cast"]:
        toc_entries.append(("cast", "Main characters"))
    for ch in data["chapters"]:
        toc_entries.append((f"ch{ch['num']:02d}", f"{ch['num']}. {ch['title']}"))
    toc_entries.append(("theend", "The End"))
    toc_entries.append(("about", "About the Author"))
    bookmark_id = insert_kindle_toc(doc, toc_entries, bookmark_id)

    if data["cast"]:
        h = doc.add_heading("Main characters", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bookmark(h, "cast", bookmark_id)
        bookmark_id += 1
        for c in data["cast"]:
            add_plain(doc, f"{c['name']} — {c['text']}", indent=0)

    for ch in data["chapters"]:
        h = doc.add_heading(f"Chapter {ch['num']}. {ch['title']}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bookmark(h, f"ch{ch['num']:02d}", bookmark_id)
        bookmark_id += 1
        first_body = [True]
        for block in ch["blocks"]:
            write_block_docx(doc, block, first_body)

    h = doc.add_heading("The End", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bookmark(h, "theend", bookmark_id)
    bookmark_id += 1
    add_plain(doc, book["end_line"], center=True, italic=True)
    if data.get("end_image"):
        add_picture(doc, data["end_image"]["path"], data["end_image"].get("caption", ""))

    h = doc.add_heading("About the Author", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bookmark(h, "about", bookmark_id)
    add_plain(doc, book["blurb"], indent=0)
    add_plain(doc, ABOUT_AUTHOR, indent=0.25)
    if book["next_note"]:
        add_plain(doc, book["next_note"], indent=0.25)

    doc.save(docx_path)


def write_listing_files(book: dict, out_dir: Path) -> None:
    if book["id"] == "home":
        desc = """Paste into KDP → Kindle eBook → Description. Amazon shows the first paragraph before “Read more,” so keep that hook as-is.

HTML is allowed: b, i, br, p, ul, li.

---

<p><b>Star is a toy sheep who wonders why she exists. Then someone chooses her.</b></p>
<p>Star is a toy sheep who wonders why she exists. When she becomes Sophia’s birthday gift, she finds friends on the toy farm — but a missing cow named Mia pulls her into her first real adventure.</p>
<p>So begins <i>Star Finds a Home</i>, a toy-farm story by young author Anshika Mahesh. Star finds her name, brings Mia home, and learns her purpose is family.</p>
<p><b>This book is for:</b></p>
<ul>
<li>Readers ages 6–8 who like talking toys and farm animals</li>
<li>Families who want a short illustrated chapter book about belonging</li>
<li>Anyone who has ever wondered where they fit</li>
</ul>
<p>Twelve chapters, with pictures. A <i>Star of the Toy Farm</i> story.</p>
<p><i>Star Finds a Home</i> · Written by Anshika Mahesh</p>
"""
        keys = """Paste one phrase into each of KDP’s 7 keyword boxes. Do not repeat the title or author name. Each is under 50 characters.

1. toy sheep finds a family kids book
2. illustrated farm toys chapter book 6 to 8
3. birthday gift toy adventure for kids
4. missing cow rescue toy farm story
5. friendship and belonging early chapter book
6. talking toys bedtime story for children
7. girl and her toy farm friends
"""
    else:
        desc = """Paste into KDP → Kindle eBook → Description. Amazon shows the first paragraph before “Read more,” so keep that hook as-is.

HTML is allowed: b, i, br, p, ul, li.

---

<p><b>A sleepover, a riddle-loving goat, and a midnight feast — Star learns bravery is looking after friends when the night gets messy.</b></p>
<p>A sleepover, a show-and-tell prize, a riddle-loving goat, and a midnight feast — Star learns that bravery isn’t only mountains. It’s looking after friends when the night gets messy.</p>
<p><i>Star and the Midnight Feast</i> is a toy-farm story by young author Anshika Mahesh. Best after <i>Star Finds a Home</i>.</p>
<p><b>This book is for:</b></p>
<ul>
<li>Readers ages 6–8 who liked Star Finds a Home</li>
<li>Families who enjoy sleepovers, talking toys, and a bit of mischief</li>
<li>Anyone who wants a story about looking after friends</li>
</ul>
<p>Twelve chapters, with pictures. A <i>Star of the Toy Farm</i> story.</p>
<p><i>Star and the Midnight Feast</i> · Written by Anshika Mahesh</p>
"""
        keys = """Paste one phrase into each of KDP’s 7 keyword boxes. Do not repeat the title or author name. Each is under 50 characters.

1. midnight feast sleepover kids book
2. riddle goat toy farm adventure
3. show and tell talking toys story
4. looking after friends early chapter book
5. toy sheep sleepover for ages 6 to 8
6. bravery and friendship farm toys
7. illustrated kids sleepover chapter book
"""
    (out_dir / "amazon-description.txt").write_text(desc, encoding="utf-8")
    (out_dir / "amazon-keywords.txt").write_text(keys, encoding="utf-8")


def write_series_files() -> None:
    series_dir = HERE / "series"
    series_dir.mkdir(parents=True, exist_ok=True)
    (series_dir / "amazon-series.txt").write_text(
        """Paste these into KDP → Bookshelf → Create series page (or Add to series).
Do not type Book 1 / Book 2 into the series title, subtitle, or this description.

============================================================
SERIES TITLE
============================================================
Star of the Toy Farm

Copy-paste exactly. Do not add the word Series — Amazon already labels it as a series.
Use this same spelling on both books. A missing space or different capital letter will split them into two series.


============================================================
READING ORDER
============================================================
Choose: Ordered

Midnight Feast is meant to be read after Star Finds a Home.

When you add each title, set:

  Content type     Main content
  Display number   1     Star Finds a Home
  Display number   2     Star and the Midnight Feast

Amazon will show those numbers on the series page. That is reading order, not a “Book 1” label on the cover.


============================================================
SERIES IMAGE
============================================================
Upload:  kdp/star-of-the-toy-farm/series/series-image.jpg
Size:    1600 × 2560 px, RGB JPEG (same as a Kindle cover)

If the form does not let you upload a file, and instead shows book-cover thumbnails:
leave it. Amazon will build the series picture from the two Kindle covers once both books are live.
Until the second book is live, shoppers may only see the first cover.


============================================================
DESCRIPTION
============================================================
Amazon shows the first paragraph before “Read more.” If you leave this blank, Amazon uses the first book’s description instead. Paste the HTML below into the series description box.

---

<p><b>A toy sheep, a farm of friends, and adventures that start after bedtime.</b></p>
<p>Star is a toy sheep who wonders why she exists. When she becomes Sophia’s birthday gift, she finds a name, a family, and friends who talk when grown-ups aren’t listening.</p>
<p>The <i>Star of the Toy Farm</i> stories follow Star, Bundle, Mia, and the farm through a missing-cow rescue, a mountain day, a sleepover, a riddle-loving goat, and a midnight feast.</p>
<p>Written by young author Anshika Mahesh. For readers ages 6–8 who like talking toys, friendship, and a little bravery. From a kid, to a kid.</p>
<p><b>Read in this order:</b></p>
<ul>
<li><i>Star Finds a Home</i></li>
<li><i>Star and the Midnight Feast</i></li>
</ul>


============================================================
AFTER YOU SAVE
============================================================
- Language: English
- Do not enroll in KDP Select until the full story is no longer free on anshikamahesh.com
- The public series page appears within about 72 hours after at least two main-content titles are Live
- Do not add an Amazon button on the website until you have the live series or product URL
""",
        encoding="utf-8",
    )
    series_book = dict(BOOKS[0])
    series_book["title_lines"] = ["STAR OF THE", "TOY FARM"]
    series_book["subtitle"] = "Stories by Anshika Mahesh"
    build_cover_jpeg(series_book, series_dir / "series-image.jpg")


def write_readme() -> None:
    (HERE / "README.md").write_text(
        """# Star of the Toy Farm — KDP Kindle files

Two Kindle titles, same series. Text is Anshika’s, not rewritten. Do **not** put Book 1 or Book 2 on the cover, subtitle, or series display name.

Rebuild after any text or art change:

```bash
python3 kdp/star-of-the-toy-farm/build_kindle.py
```

## Files to upload

| Title | Folder | Manuscript | Cover |
| --- | --- | --- | --- |
| **Star Finds a Home** | `star-finds-a-home/` | `star-finds-a-home-kindle.docx` | `kindle-cover.jpg` |
| **Star and the Midnight Feast** | `star-and-the-midnight-feast/` | `star-and-the-midnight-feast-kindle.docx` | `kindle-cover.jpg` |

If KDP rejects the `.docx`, use the `.epub` in the same folder. Do **not** upload the website `.html`.

## KDP Kindle eBook settings (each title)

Create **two** Kindle eBooks on the parent KDP account (account holder must be 18+). Author name stays **Anshika Mahesh**.

1. Language: English.
2. Title: use the table above. Subtitle: **A Star of the Toy Farm story**. Do not add Book 1 / Book 2.
3. Series: name the series **Star of the Toy Farm**. Sequence 1 = Star Finds a Home, sequence 2 = Star and the Midnight Feast. The public titles stay the adventure names.
4. Age: **6–8**. Categories: Juvenile Fiction / Animals / Toys, Dolls & Puppets; Social Themes / Friendship.
5. Description and keywords: paste from that folder’s `amazon-description.txt` and `amazon-keywords.txt`.
6. Upload the `.docx`, then `kindle-cover.jpg` (1600 × 2560, RGB JPEG). Re-upload the `.docx` after a rebuild so Kindle can see the Contents list (Heading 1 chapters + bookmark `toc`).
7. Tick the **AI-generated images** disclosure.
8. Preview on phone, tablet, and e-ink in KDP Previewer.
9. Primary marketplace: **Amazon.in**. Set **₹149** there, and **$2.99** on Amazon.com. The book still sells worldwide.
10. India pays 70% royalty only if the book is in KDP Select (price band ₹99–₹599). Without Select, Amazon.in is 35%. US/UK can still be 70% at $2.99 without Select.
11. Do not enroll in KDP Select until the full story is no longer free on anshikamahesh.com.
12. Publish Star Finds a Home first. Link Midnight Feast as the next book in the same series.

## Series page (Bookshelf → Create series page)

Paste from `series/amazon-series.txt`. Upload `series/series-image.jpg` if the form asks for a series image.

| Field | Enter |
| --- | --- |
| Series title | **Star of the Toy Farm** (do not add the word Series) |
| Reading order | **Ordered** — 1 = Star Finds a Home, 2 = Star and the Midnight Feast |
| Series image | `series/series-image.jpg` (1600 × 2560). If the form only offers cover thumbnails, leave it. |
| Description | Paste the HTML in `series/amazon-series.txt`. Do not leave blank if you want a series-wide blurb. |

Paperback interiors are not in this folder yet. Kindle can go live on its own.

Do not add the Amazon button on anshikamahesh.com until a title is **Live** and you have the product URL / ASIN.
""",
        encoding="utf-8",
    )


def build_one(book: dict) -> None:
    out_dir = HERE / book["out_dir"]
    out_dir.mkdir(parents=True, exist_ok=True)
    cover_path = out_dir / "kindle-cover.jpg"
    docx_path = out_dir / f"{book['slug']}-kindle.docx"
    epub_path = out_dir / f"{book['slug']}-kindle.epub"

    data = parse_story(book)
    if len(data["chapters"]) != book["expected_chapters"]:
        raise SystemExit(
            f"{book['title']}: expected {book['expected_chapters']} chapters, got {len(data['chapters'])}"
        )
    missing = [str(book["cover_src"])] if not book["cover_src"].exists() else []
    for ch in data["chapters"]:
        for block in chapter_images(ch):
            if not block["path"].exists():
                missing.append(str(block["path"]))
    if missing:
        raise SystemExit("missing images:\n" + "\n".join(missing))

    book = dict(book)
    book["book_id"] = str(uuid.uuid5(uuid.NAMESPACE_URL, book["uuid_url"]))
    build_cover_jpeg(book, cover_path)
    build_docx(book, data, docx_path)
    size = build_epub(book, data, cover_path, epub_path)
    write_listing_files(book, out_dir)
    n_img = sum(len(chapter_images(ch)) for ch in data["chapters"]) + (1 if data.get("end_image") else 0)
    print(f"{book['title']}")
    print(f"  chapters {len(data['chapters'])}  pictures {n_img}")
    print(f"  {cover_path.relative_to(ROOT)}  {COVER_W}×{COVER_H}")
    print(f"  {docx_path.relative_to(ROOT)}  {docx_path.stat().st_size / (1024 * 1024):.2f} MB")
    print(f"  {epub_path.relative_to(ROOT)}  {size / (1024 * 1024):.2f} MB")


def main() -> None:
    write_readme()
    write_series_files()
    for book in BOOKS:
        build_one(book)


if __name__ == "__main__":
    main()
