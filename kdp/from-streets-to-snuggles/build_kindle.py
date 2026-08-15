#!/usr/bin/env python3
"""Build Kindle files for From Streets to Snuggles.

Produces:
  from-streets-to-snuggles-kindle.docx   ← upload this to KDP first
  from-streets-to-snuggles-kindle.epub   ← EPUB 2 fallback
  kindle-cover.jpg

Source text is copied from the site HTML without rewriting Anshika's story.

    python3 kdp/from-streets-to-snuggles/build_kindle.py
"""
from __future__ import annotations

import html as html_lib
import io
import re
import sys
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageEnhance, ImageFont

sys.path.insert(0, str(Path(__file__).resolve().parent))
from build import COVER_SRC, FONT_DIR, HERE, PLATE_DIR, STORY, parse_story

TITLE = "From Streets to Snuggles"
SUBTITLE = "A tale of kindness, rescue, and belonging"
AUTHOR = "Anshika Mahesh"
LANG = "en-US"
COVER_W, COVER_H = 1600, 2560
PLATE_MAX_W = 1200

COVER_PATH = HERE / "kindle-cover.jpg"
EPUB_PATH = HERE / "from-streets-to-snuggles-kindle.epub"
DOCX_PATH = HERE / "from-streets-to-snuggles-kindle.docx"

BLURB = (
    "In a heartwarming tale of love, friendship, and second chances, "
    "Fluffy and Snowy find their way from the harsh streets to a loving home "
    "— thanks to one kind boy named Sam."
)


def esc(text: str) -> str:
    return html_lib.escape(text, quote=True)


def chapter_alts() -> dict[int, str]:
    raw = STORY.read_text(encoding="utf-8")
    alts = {}
    for num, body in re.findall(
        r'<section class="page chapter" id="ch(\d+)">(.*?)</section>',
        raw,
        re.S,
    ):
        m = re.search(r'<img[^>]+alt="([^"]*)"', body)
        alts[int(num)] = html_lib.unescape(m.group(1)) if m else f"Chapter {int(num)} illustration"
    return alts


def jpeg_bytes(src: Path, max_w: int, quality: int = 82) -> bytes:
    im = Image.open(src).convert("RGB")
    if im.width > max_w:
        h = max(1, int(im.height * max_w / im.width))
        im = im.resize((max_w, h), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    # Baseline JPEG only — progressive files break some Kindle converters.
    im.save(buf, "JPEG", quality=quality, optimize=True, progressive=False)
    return buf.getvalue()


def build_cover_jpeg() -> None:
    img = Image.new("RGB", (COVER_W, COVER_H), "#1c2430")
    art = ImageEnhance.Contrast(Image.open(COVER_SRC).convert("RGB")).enhance(1.02)
    scale = max(COVER_W / art.width, COVER_H / art.height)
    nw, nh = int(art.width * scale), int(art.height * scale)
    art = art.resize((nw, nh), Image.Resampling.LANCZOS)
    img.paste(art, ((COVER_W - nw) // 2, (COVER_H - nh) // 2))

    overlay = Image.new("RGBA", (COVER_W, COVER_H), (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)
    od.rectangle((0, 0, COVER_W, 430), fill=(28, 36, 48, 160))
    od.rectangle((0, COVER_H - 280, COVER_W, COVER_H), fill=(28, 36, 48, 160))
    img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype(str(FONT_DIR / "Georgia Bold.ttf"), 92)
        sub_font = ImageFont.truetype(str(FONT_DIR / "Georgia Italic.ttf"), 36)
        author_font = ImageFont.truetype(str(FONT_DIR / "Georgia Bold.ttf"), 34)
    except OSError:
        title_font = sub_font = author_font = ImageFont.load_default()

    def centre(text: str, font, y: float, fill="white") -> None:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(((COVER_W - tw) / 2, y - th / 2), text, font=font, fill=fill)

    centre("FROM STREETS", title_font, 150)
    centre("TO SNUGGLES", title_font, 250)
    centre(SUBTITLE, sub_font, 340)
    centre("ANSHIKA MAHESH", author_font, COVER_H - 140)
    img.save(COVER_PATH, "JPEG", quality=90, optimize=True, progressive=False, dpi=(300, 300))


def xhtml(title: str, body: str) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" '
        '"http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">\n'
        f'<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{LANG}">\n'
        "<head>\n"
        f"<title>{esc(title)}</title>\n"
        '<meta http-equiv="Content-Type" content="application/xhtml+xml; charset=utf-8" />\n'
        '<link rel="stylesheet" type="text/css" href="styles.css" />\n'
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )


def styles_css() -> str:
    return """body {
  font-family: Georgia, "Times New Roman", serif;
  font-size: 1em;
  line-height: 1.4;
  margin: 0;
  text-align: justify;
}
h1 {
  font-size: 1.4em;
  text-align: center;
  margin: 0.4em 0 0.8em 0;
  page-break-after: avoid;
}
.kicker {
  text-align: center;
  font-size: 0.85em;
  font-weight: bold;
  text-transform: uppercase;
  margin: 1.2em 0 0.3em 0;
}
.center { text-align: center; text-indent: 0; }
p { margin: 0 0 0.6em 0; text-indent: 1.2em; }
p.first { text-indent: 0; }
img.plate {
  width: 100%;
  height: auto;
  display: block;
  margin: 0.5em 0 1em 0;
}
"""


def container_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
"""


def content_opf(book_id: str, chapters: list[dict]) -> str:
    items = [
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="css" href="styles.css" media-type="text/css"/>',
        '<item id="cover-image" href="images/cover.jpg" media-type="image/jpeg"/>',
        '<item id="titlepage" href="title.xhtml" media-type="application/xhtml+xml"/>',
        '<item id="copyright" href="copyright.xhtml" media-type="application/xhtml+xml"/>',
        '<item id="contents" href="contents.xhtml" media-type="application/xhtml+xml"/>',
    ]
    spine = [
        '<itemref idref="titlepage"/>',
        '<itemref idref="copyright"/>',
        '<itemref idref="contents"/>',
    ]
    for ch in chapters:
        n = f"{ch['num']:02d}"
        items.append(f'<item id="ch{n}" href="ch{n}.xhtml" media-type="application/xhtml+xml"/>')
        items.append(f'<item id="img{n}" href="images/ch-{n}.jpg" media-type="image/jpeg"/>')
        spine.append(f'<itemref idref="ch{n}"/>')
    items.append('<item id="theend" href="end.xhtml" media-type="application/xhtml+xml"/>')
    items.append('<item id="about" href="about.xhtml" media-type="application/xhtml+xml"/>')
    spine.append('<itemref idref="theend"/>')
    spine.append('<itemref idref="about"/>')
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
    <dc:identifier id="BookId" opf:scheme="UUID">urn:uuid:{book_id}</dc:identifier>
    <dc:title>{esc(TITLE)}</dc:title>
    <dc:creator opf:role="aut">{esc(AUTHOR)}</dc:creator>
    <dc:language>{LANG}</dc:language>
    <dc:publisher>{esc(AUTHOR)}</dc:publisher>
    <dc:rights>Copyright &#169; 2025 {esc(AUTHOR)}. All rights reserved.</dc:rights>
    <dc:description>When Sam rescues two street dogs, their bond grows through baths, the vet, new friends, and a holiday in Goa.</dc:description>
    <meta name="cover" content="cover-image"/>
  </metadata>
  <manifest>
    {chr(10).join("    " + i for i in items)}
  </manifest>
  <spine toc="ncx">
    {chr(10).join("    " + s for s in spine)}
  </spine>
  <guide>
    <reference type="cover" title="Cover" href="images/cover.jpg"/>
    <reference type="toc" title="Contents" href="contents.xhtml"/>
    <reference type="text" title="Start Reading" href="ch01.xhtml"/>
  </guide>
</package>
"""


def toc_ncx(book_id: str, chapters: list[dict]) -> str:
    points = [
        '<navPoint id="nav1" playOrder="1"><navLabel><text>Title</text></navLabel><content src="title.xhtml"/></navPoint>',
        '<navPoint id="nav2" playOrder="2"><navLabel><text>Contents</text></navLabel><content src="contents.xhtml"/></navPoint>',
    ]
    order = 3
    for ch in chapters:
        n = f"{ch['num']:02d}"
        points.append(
            f'<navPoint id="nav-ch{n}" playOrder="{order}">'
            f"<navLabel><text>Chapter {ch['num']}. {esc(ch['title'])}</text></navLabel>"
            f'<content src="ch{n}.xhtml"/></navPoint>'
        )
        order += 1
    points.append(
        f'<navPoint id="nav-end" playOrder="{order}"><navLabel><text>The End</text></navLabel><content src="end.xhtml"/></navPoint>'
    )
    points.append(
        f'<navPoint id="nav-about" playOrder="{order + 1}"><navLabel><text>About the Book</text></navLabel><content src="about.xhtml"/></navPoint>'
    )
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{book_id}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{esc(TITLE)}</text></docTitle>
  <navMap>
    {chr(10).join("    " + p for p in points)}
  </navMap>
</ncx>
"""


def title_page() -> str:
    return xhtml(
        TITLE,
        f'<p class="kicker">A Kindle edition</p>\n'
        f"<h1>{esc(TITLE)}</h1>\n"
        f'<p class="center">{esc(SUBTITLE)}</p>\n'
        f'<p class="center">{esc(AUTHOR)}</p>\n'
        f'<p class="first center">{esc(BLURB)}</p>',
    )


def copyright_page() -> str:
    return xhtml(
        "Copyright",
        f"<h1>Copyright</h1>\n"
        f'<p class="first"><strong>{esc(TITLE)}</strong></p>\n'
        f'<p class="first">Copyright &#169; 2025 {esc(AUTHOR)}</p>\n'
        f'<p class="first">First Kindle edition 2026</p>\n'
        f'<p class="first">All rights reserved. No part of this book may be reproduced without permission from the author, except for brief quotations in a review.</p>\n'
        f'<p class="first">This is a work of fiction. Names, characters, and incidents are from the author&#8217;s imagination.</p>\n'
        f'<p class="first">Also available as a paperback. Read free at anshikamahesh.com.</p>',
    )


def contents_page(toc: list[dict]) -> str:
    items = []
    for i, row in enumerate(toc, 1):
        items.append(f'<p class="first"><a href="ch{i:02d}.xhtml">{i}. {esc(row["title"])}</a></p>')
    return xhtml("Contents", '<h1 id="toc">Contents</h1>\n' + "\n".join(items))


def chapter_page(ch: dict, alt: str) -> str:
    n = f"{ch['num']:02d}"
    paras = []
    for i, text in enumerate(ch["paragraphs"]):
        cls = ' class="first"' if i == 0 else ""
        paras.append(f"<p{cls}>{esc(text)}</p>")
    return xhtml(
        f"Chapter {ch['num']}. {ch['title']}",
        f'<p class="kicker">Chapter {ch["num"]}</p>\n'
        f"<h1>{esc(ch['title'])}</h1>\n"
        f'<img class="plate" src="images/ch-{n}.jpg" alt="{esc(alt)}" />\n'
        + "\n".join(paras),
    )


def end_page() -> str:
    return xhtml(
        "The End",
        '<h1>The End</h1>\n<p class="first center">And now we leave them to dream of new adventures.</p>',
    )


def about_page() -> str:
    return xhtml(
        "About the Book",
        "<h1>About the Book</h1>\n"
        '<p class="first">When Sam, a kind-hearted young boy, rescues Fluffy, their bond grows stronger with each passing day. They embark on countless adventures&#8212;learning tricks, visiting the vet, making new friends, and even taking a holiday to Goa!</p>\n'
        '<p>Perfect for animal lovers and young readers, this touching story reminds us that sometimes, the greatest journeys begin with a single act of kindness.</p>\n'
        '<p>Anshika Mahesh writes stories about friendship, courage, and kindness. Read more at anshikamahesh.com.</p>',
    )


def _zipinfo(name: str, stored: bool = False) -> zipfile.ZipInfo:
    info = zipfile.ZipInfo(name, date_time=(2026, 8, 15, 0, 0, 0))
    info.compress_type = zipfile.ZIP_STORED if stored else zipfile.ZIP_DEFLATED
    info.create_system = 3
    info.external_attr = 0o644 << 16
    info.extra = b""
    return info


def build_epub(data: dict) -> int:
    alts = chapter_alts()
    book_id = str(uuid.uuid5(uuid.NAMESPACE_URL, "https://anshikamahesh.com/stories/from-streets-to-snuggles"))
    files: list[tuple[zipfile.ZipInfo, bytes]] = []

    def add(name: str, content: str | bytes, stored: bool = False) -> None:
        payload = content.encode("utf-8") if isinstance(content, str) else content
        files.append((_zipinfo(name, stored=stored), payload))

    add("mimetype", "application/epub+zip", stored=True)
    add("META-INF/container.xml", container_xml())
    add("OEBPS/content.opf", content_opf(book_id, data["chapters"]))
    add("OEBPS/toc.ncx", toc_ncx(book_id, data["chapters"]))
    add("OEBPS/styles.css", styles_css())
    add("OEBPS/title.xhtml", title_page())
    add("OEBPS/copyright.xhtml", copyright_page())
    add("OEBPS/contents.xhtml", contents_page(data["toc"]))
    add("OEBPS/end.xhtml", end_page())
    add("OEBPS/about.xhtml", about_page())
    add("OEBPS/images/cover.jpg", COVER_PATH.read_bytes())

    missing = []
    for ch in data["chapters"]:
        n = f"{ch['num']:02d}"
        plate = PLATE_DIR / f"ch-{n}.png"
        if not plate.exists():
            missing.append(str(plate))
            continue
        add(f"OEBPS/ch{n}.xhtml", chapter_page(ch, alts.get(ch["num"], f"Chapter {ch['num']} illustration")))
        add(f"OEBPS/images/ch-{n}.jpg", jpeg_bytes(plate, PLATE_MAX_W))
    if missing:
        raise SystemExit("missing plates:\n" + "\n".join(missing))

    with zipfile.ZipFile(EPUB_PATH, "w") as zf:
        for info, payload in files:
            zf.writestr(info, payload)
    return EPUB_PATH.stat().st_size


def _set_run_font(run, name: str = "Georgia", size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)


def _center(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)


def _body(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15


def _bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _page_break(paragraph) -> None:
    paragraph.paragraph_format.page_break_before = True


def _add_toc_link(paragraph, anchor: str, label: str) -> None:
    """Visible chapter link. Kindle keeps these; it often drops Word TOC fields."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Georgia")
    fonts.set(qn("w:hAnsi"), "Georgia")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "3D7F99")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_docx(data: dict) -> None:
    doc = Document()
    core = doc.core_properties
    core.title = TITLE
    core.author = AUTHOR
    core.subject = SUBTITLE
    core.category = "Juvenile Fiction"

    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)
    heading = styles["Heading 1"]
    heading.font.name = "Georgia"
    heading.font.size = Pt(18)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.page_break_before = True
    heading_ppr = heading.element.get_or_add_pPr()
    if heading_ppr.find(qn("w:outlineLvl")) is None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        heading_ppr.append(outline)

    p = doc.add_paragraph()
    _center(p)
    run = p.add_run(TITLE)
    _set_run_font(run, size=28, bold=True)
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run(SUBTITLE)
    _set_run_font(run, size=13, italic=True)
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run(AUTHOR)
    _set_run_font(run, size=12, bold=True)
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run(BLURB)
    _set_run_font(run, size=12, italic=True)

    p = doc.add_paragraph()
    _center(p)
    _page_break(p)
    run = p.add_run("Copyright")
    _set_run_font(run, size=18, bold=True)
    for line in (
        f"{TITLE}",
        f"Copyright © 2025 {AUTHOR}",
        "First Kindle edition 2026",
        "All rights reserved. No part of this book may be reproduced without permission from the author, except for brief quotations in a review.",
        "This is a work of fiction. Names, characters, and incidents are from the author’s imagination.",
        "Also available as a paperback. Read free at anshikamahesh.com.",
    ):
        p = doc.add_paragraph()
        _body(p)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(line)
        _set_run_font(run, size=11)

    toc_title = doc.add_paragraph()
    _center(toc_title)
    _page_break(toc_title)
    run = toc_title.add_run("Contents")
    _set_run_font(run, size=18, bold=True)
    _bookmark(toc_title, "toc", 1)

    bookmark_id = 2
    for ch in data["chapters"]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(8)
        _add_toc_link(p, f"ch{ch['num']:02d}", f"{ch['num']}. {ch['title']}")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(8)
    _add_toc_link(p, "theend", "The End")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    _add_toc_link(p, "about", "About the Book")

    for ch in data["chapters"]:
        h = doc.add_heading(f"Chapter {ch['num']}. {ch['title']}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bookmark(h, f"ch{ch['num']:02d}", bookmark_id)
        bookmark_id += 1
        plate = PLATE_DIR / f"ch-{ch['num']:02d}.png"
        if plate.exists():
            pic = doc.add_paragraph()
            _center(pic)
            pic.paragraph_format.space_after = Pt(12)
            run = pic.add_run()
            run.add_picture(io.BytesIO(jpeg_bytes(plate, PLATE_MAX_W)), width=Inches(4.5))
        for i, text in enumerate(ch["paragraphs"]):
            p = doc.add_paragraph()
            _body(p)
            p.paragraph_format.first_line_indent = Inches(0 if i == 0 else 0.25)
            run = p.add_run(text)
            _set_run_font(run, size=12)

    h = doc.add_heading("The End", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bookmark(h, "theend", bookmark_id)
    bookmark_id += 1
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run("And now we leave them to dream of new adventures.")
    _set_run_font(run, size=12, italic=True)

    h = doc.add_heading("About the Book", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bookmark(h, "about", bookmark_id)
    abouts = [
        "When Sam, a kind-hearted young boy, rescues Fluffy, their bond grows stronger with each passing day. They embark on countless adventures—learning tricks, visiting the vet, making new friends, and even taking a holiday to Goa!",
        "Perfect for animal lovers and young readers, this touching story reminds us that sometimes, the greatest journeys begin with a single act of kindness.",
        "Anshika Mahesh writes stories about friendship, courage, and kindness. Read more at anshikamahesh.com.",
    ]
    for i, text in enumerate(abouts):
        p = doc.add_paragraph()
        _body(p)
        p.paragraph_format.first_line_indent = Inches(0 if i == 0 else 0.25)
        run = p.add_run(text)
        _set_run_font(run, size=12)

    doc.save(DOCX_PATH)


def main() -> None:
    data = parse_story()
    if len(data["chapters"]) != 25:
        raise SystemExit(f"expected 25 chapters, got {len(data['chapters'])}")
    if len(data["toc"]) != 25:
        raise SystemExit(f"expected 25 contents rows, got {len(data['toc'])}")
    if not COVER_SRC.exists():
        raise SystemExit(f"missing cover: {COVER_SRC}")
    build_cover_jpeg()
    build_docx(data)
    size = build_epub(data)
    print(f"kindle-cover.jpg  {COVER_W}×{COVER_H}")
    print(f"{DOCX_PATH.name}  {DOCX_PATH.stat().st_size / (1024 * 1024):.2f} MB")
    print(f"{EPUB_PATH.name}  {size / (1024 * 1024):.2f} MB")


if __name__ == "__main__":
    main()
